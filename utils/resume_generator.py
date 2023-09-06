from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re
import utils.prompts as pr

def create_doc():
    # Create a new document
    doc = Document()
    section = doc.sections[0]
    # Setting margins for all sides to 0.5 inches
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    margin = section.page_width - section.right_margin - section.left_margin
    return doc, margin
    # Accessing the default section of the document          

doc = create_doc()

# Define the color for "blue-gray, accent 3, 25% darker" (This is an example color. Adjust it!)
blue_gray = RGBColor(64, 64, 64) # Example color, adjust accordingly

def add_bottom_border_to_paragraph(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    
    bottom = (
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:bottom w:val="single" w:sz="24" w:space="1" w:color="000000"/>'
        '</w:pBdr>'
    )
    
    pPr.insert(0, parse_xml(bottom))



def add_custom_paragraph(doc,margin,text1, text2, font_color=None, bold=False, add_line=False, font_size1=16, font_size2=10):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(margin, WD_TAB_ALIGNMENT.RIGHT)
    
    if add_line:
        add_bottom_border_to_paragraph(p)

    run1 = p.add_run(text1)
    run1.font.name = 'Helvita Neue Light'
    run1.font.size = Pt(font_size1)
    if font_color:
        run1.font.color.rgb = font_color
    if bold:
        run1.font.bold = True

    p.add_run('\t')

    run2 = p.add_run(text2)
    run2.font.name = 'Helvita Neue Light'
    run2.font.size = Pt(font_size2)
    return p

def add_title_section(doc,title_text,color):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title_text)
    run.font.name = 'Avenir Book'
    run.font.color.rgb = color
    run.font.size = Pt(14)
    run.bold = True
    return p

def add_career_summary(doc,summary_text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Set the alignment to Justify
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0  # Set line spacing to 1.0
    p.paragraph_format.line_spacing_rule = 1  # Use single spacing
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(summary_text)
    run.font.name = 'Helvetica Neue Light'
    run.font.size = Pt(10)
    return p


def add_key_accomplishments_sections(doc,items,section_title_color):
    # Add title
    title = doc.add_paragraph('Key Accomplishments')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.font.name = 'Avenir Light'
    run.font.size = Pt(14)
    run.font.color.rgb = section_title_color
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.space_before = Pt(3)

    # Add top border to title using low-level XML element
    border_xml = r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="1" w:color="auto"/></w:pBdr>'
    title._element.get_or_add_pPr().insert(0, parse_xml(border_xml))

    # Role Accomplishments/Duties
    for accomplishements in items:
        bullet = doc.add_paragraph()
        
        # Create square bullet
        run = bullet.add_run("\u25A0")
        run.font.name = 'Helvetica Neue Light'
        run.font.size = Pt(10)
        
        # Add duty text
        run = bullet.add_run(f""" {accomplishements}""")
        run.font.name = 'Helvetica Neue Light'
        run.font.size = Pt(10)
        
        # Adjust bullet point formatting
        bullet_format = bullet.paragraph_format
        bullet_format.left_indent = Inches(0.15)
        bullet_format.first_line_indent = Inches(-0.15)
        bullet_format.space_after = Pt(2)
        # Set line spacing to 1.0
        bullet_format.line_spacing = 1.0


def add_key_competencies_section(doc,items,section_title_color):
    # Check if the items list has 12 entries
    if len(items) != 12:
        print("Warning: The items list should have 12 entries!")
        return
    
    # Add title
    title = doc.add_paragraph('Key Competencies')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.font.name = 'Avenir Light'
    run.font.size = Pt(14)
    run.font.color.rgb = section_title_color
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.space_before = Pt(3)
    
    # Add top border to title using low-level XML element
    border_xml = r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="1" w:color="auto"/></w:pBdr>'
    title._element.get_or_add_pPr().insert(0, parse_xml(border_xml))
    
    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False

    # Split the items list into 3 columns with 4 items each
    col1_items = items[:4]
    col2_items = items[4:8]
    col3_items = items[8:]

    # Populate the table cells
    table.cell(0, 0).text = '\n'.join(col1_items)
    table.cell(0, 1).text = '\n'.join(col2_items)
    table.cell(0, 2).text = '\n'.join(col3_items)

    # Style the cell content
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Helvetica Neue Light'
                    run.font.size = Pt(10)
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_experience_header(doc,section_title_color):
    title = doc.add_paragraph()
    run = title.add_run("Professional Experience")
    run.font.name = 'Avenir Light'
    run.font.size = Pt(14)
    run.font.color.rgb = section_title_color
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.space_before = Pt(3)
    
    # Add top border to title using low-level XML element
    border_xml = r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="1" w:color="auto"/></w:pBdr>'
    title._element.get_or_add_pPr().insert(0, parse_xml(border_xml))

def add_career_section(doc, experience,job_description,bullets,company_color):
    
    # Company Details
    p = doc.add_paragraph()
    run = p.add_run(f"{experience['company_name']} | {experience['location']} • {experience['years']}")
    run.bold = True
    run.font.name = 'Helvitca Neue Medium Italic'
    run.font.size = Pt(11)
    run.font.color.rgb = company_color
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)

    # Optional Company Summary
    if 'summary' in experience:
        summary_p = doc.add_paragraph()
        summary_run = summary_p.add_run(experience['summary'])
        summary_run.font.name = 'Helvetica Neue Light'
        summary_run.font.size = Pt(10)
        summary_p.paragraph_format.space_after = Pt(2)
    
    for role in experience["roles"]:
        # Role Title and Dates
        p = doc.add_paragraph()
        run = p.add_run(f"{role['title']} | {role['dates']}")
        run.bold = True
        run.font.name = 'Helvetica Neue Light'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)
    
        selected_accomplishments = pr.response(pr.select_bullets(role["accomplishments"],job_description,bullets))
        # Role Accomplishments/Duties
        for accomplishment in selected_accomplishments:
            bullet = doc.add_paragraph()
            
            # Create square bullet
            run = bullet.add_run("\u25A0")
            run.font.name = 'Helvetica Neue Light'
            run.font.size = Pt(10)
            
            # Add duty text
            run = bullet.add_run(f""" {accomplishment}""")
            run.font.name = 'Helvetica Neue Light'
            run.font.size = Pt(10)
            
            # Adjust bullet point formatting
            bullet_format = bullet.paragraph_format
            bullet_format.left_indent = Inches(0.15)
            bullet_format.first_line_indent = Inches(-0.15)
            bullet_format.space_after = Pt(2)
            # Set line spacing to 1.0
            bullet_format.line_spacing = 1.0
            #bullet_format.space_after = Pt(2)

    # Add a spacing of 6pt after the career section
    space_after_section = doc.add_paragraph()
    space_after_section.paragraph_format.space_after = Pt(6)


def add_education_and_training_section(doc,educations,section_title_color,education_color):
        # Title: "Education & Training"
        title = doc.add_paragraph()
        run = title.add_run("Education & Training")
        run.font.name = 'Avenir Light'
        run.font.size = Pt(14)
        run.font.color.rgb = section_title_color
        title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.paragraph_format.space_after = Pt(6)
        title.paragraph_format.space_before = Pt(3)

        # Add top border to title using low-level XML element
        border_xml = r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="1" w:color="auto"/></w:pBdr>'
        title._element.get_or_add_pPr().insert(0, parse_xml(border_xml))

        # Define the color
        # Add each education/training item
        for education in educations:
            item = doc.add_paragraph()
            degree, institution = education.split(', ', 1)  # Split the education into degree and institution

            # Add degree (in black)
            run = item.add_run(degree)
            run.bold = True
            run.font.name = 'Helvetica Neue Light'
            run.font.size = Pt(11)
            item.add_run(", ")

            # Add institution (in blue)
            run = item.add_run(institution)
            run.bold = True
            run.font.color.rgb = education_color
            run.font.name = 'Helvetica Neue Light'
            item.paragraph_format.space_after = Pt(6)

def cover_letter_body(doc,summary_text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Set the alignment to Justify
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0  # Set line spacing to 1.0
    p.paragraph_format.line_spacing_rule = 1  # Use single spacing
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(summary_text)
    run.font.name = 'Helvetica Neue Light'
    run.font.size = Pt(11)
    return p

import os

def save_document_in_folder(base_path, folder_name, doc_name, doc: Document):
    # Construct the full path for the folder
    full_folder_path = os.path.join(base_path, folder_name)
    
    # Check if folder exists
    if not os.path.exists(full_folder_path):
        os.makedirs(full_folder_path)

    # Save the document in the folder
    doc_path = os.path.join(full_folder_path, doc_name)
    doc.save(doc_path)


def convert_to_list(questions_str):
    # Check if the string is in list format
    if questions_str.startswith("[") and questions_str.endswith("]"):
        try:
            # Try to evaluate the string as a list
            questions = eval(questions_str)
            if all(isinstance(q, str) for q in questions):
                return questions
        except:
            pass  # If evaluation fails, we will handle the string as a normal text

    # If not in list format, try to split based on the pattern of numbered questions
    questions = re.split(r'\d+\.\s+', questions_str)
    
    # Remove any empty strings from the list
    questions = [q.strip() for q in questions if q.strip()]

    return questions

def write_dict_to_word(data,tell_me_about_yourself=None):
    doc = Document()
    
    # Default response to the question "tell me about yourself"
    default_response = {
        "Tell me about yourself": tell_me_about_yourself
    }
    
    data.append(default_response)

    for entry in data:
        for question, response in entry.items():
            doc.add_heading(question, level=1)   # Add the question as a heading
            doc.add_paragraph(response)          # Add the response as a paragraph

    return doc


